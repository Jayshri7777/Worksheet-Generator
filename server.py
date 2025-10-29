import os
import google.generativeai as genai
# --- IMPORTS for DB and Auth ---
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from flask import render_template, redirect, url_for, flash, session # Added session
import re # For email, phone, and password validation
# --- Authlib Imports ---
from authlib.integrations.flask_client import OAuth
# -----------------------------
from flask import Flask, request, jsonify, send_file, send_from_directory
from datetime import datetime
from fpdf import FPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from flask_cors import CORS
# --- Optional Import for Image Generation ---
try:
    from pdf2image import convert_from_path
except ImportError:
    print("--- WARNING ---")
    print("pdf2image library not found. PNG/JPG export will fail.")
    print("Run 'pip install pdf2image Authlib' to fix this.") # Added Authlib here
    print("You also MUST install 'poppler' on your computer.")
    print("Windows: Download from GitHub releases or use 'choco install poppler'. Ensure Poppler's bin folder is in PATH.")
    print("macOS: 'brew install poppler'.")
    print("Linux: Use your package manager (e.g., 'apt install poppler-utils').")
    print("-----------------")
    convert_from_path = None # Set to None if import fails

# --- Flask App Setup ---
app = Flask(__name__,
            static_folder='static',    # Folder for CSS, JS
            template_folder='templates') # Folder for HTML
CORS(app) # Enable CORS if needed

# --- CONFIG ---
# IMPORTANT: Change this secret key to something truly random and secret in production!
app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY', 'default-super-secret-key-change-me-immediately')
app.config['GOOGLE_CLIENT_ID'] = os.environ.get('GOOGLE_CLIENT_ID')
app.config['GOOGLE_CLIENT_SECRET'] = os.environ.get('GOOGLE_CLIENT_SECRET')
# Database file will be created in the same directory as server.py
db_path = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'users.db')
app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{db_path}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
# OAuth server metadata URLs
app.config['GOOGLE_CONF_URL'] = 'https://accounts.google.com/.well-known/openid-configuration'
# --------------------

# --- DATABASE SETUP ---
db = SQLAlchemy(app)

# --- USER MODEL (with Name, Grade, Phone) ---
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False) # Added Name, required
    email = db.Column(db.String(120), unique=True, nullable=False) # Email remains unique & required
    phone_number = db.Column(db.String(30), unique=True, nullable=True) # Phone optional for OAuth
    grade = db.Column(db.String(50), nullable=True) # Added Grade, optional
    password_hash = db.Column(db.String(256), nullable=True) # Password optional for OAuth

    def set_password(self, password):
        # Use a strong hashing method with salt
        self.password_hash = generate_password_hash(password, method='pbkdf2:sha256', salt_length=16)

    def check_password(self, password):
        # Ensure password_hash exists before checking (for OAuth users)
        return self.password_hash and check_password_hash(self.password_hash, password)

    def __repr__(self):
        # Useful representation for debugging
        return f'<User {self.email}>'
# -------------------------

# --- LOGIN MANAGER SETUP ---
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login' # The function name (endpoint) for the login page
login_manager.login_message = "Please log in to access this page." # Message shown when login is required
login_manager.login_message_category = "warning" # Category for flash message styling

@login_manager.user_loader
def load_user(user_id):
    # Flask-Login uses this function to reload the user object from the user ID stored in the session
    # Uses .get() which is efficient for primary key lookups
    return User.query.get(int(user_id))
# -----------------------------

# --- OAuth Setup ---
oauth = OAuth(app)

# Configure Google OAuth
oauth.register(
    name='google',
    client_id=app.config['GOOGLE_CLIENT_ID'],
    client_secret=app.config['GOOGLE_CLIENT_SECRET'],
    server_metadata_url=app.config['GOOGLE_CONF_URL'],
    client_kwargs={
        'scope': 'openid email profile' # Standard scopes
    }
)
# ------------------------

# --- PDF Generation Class (Unchanged) ---
class CustomPDF(FPDF):
    # ... (class code remains the same as before) ...
    def __init__(self, title, sub_info):
        super().__init__()
        self.worksheet_title = title
        self.sub_info = sub_info
    def header(self):
        self.set_font("Arial", "B", 16)
        self.cell(0, 10, self.worksheet_title, 0, 1, "C")
        self.set_font("Arial", "", 10)
        sub_header_text = f"Date: {self.sub_info['date']}   |   Marks: {self.sub_info['marks']}   |   Topic: {self.sub_info['sub-title']}"
        self.cell(0, 10, sub_header_text, 0, 1, "C")
        self.ln(10) # Line break
    def footer(self):
        self.set_y(-15) # Position 1.5 cm from bottom
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C") # Page number centered
        self.cell(0, 10, "Generated by AI Worksheet Tool", 0, 0, "R") # Right aligned text


# --- File Creation Functions ---
def create_pdf(content, title, sub_title_info):
    # ... (function code remains the same as before) ...
    print("Generating PDF...")
    file_path = "worksheet.pdf" # Output filename
    pdf = CustomPDF(title, sub_title_info)
    pdf.add_page()
    pdf.set_font("Arial", "", 12)
    # Handle potential encoding issues when writing to PDF
    try:
        pdf.multi_cell(0, 10, content) # Width 0 means full width, height 10 per line
    except UnicodeEncodeError:
        print("Warning: Encoding issue detected in PDF generation. Using latin-1 replacement.")
        pdf.multi_cell(0, 10, content.encode('latin-1', 'replace').decode('latin-1'))
    pdf.output(file_path) # Save the PDF
    return file_path

def create_docx(content, title, sub_title_info):
    # ... (function code remains the same as before) ...
    print("Generating DOCX...")
    file_path = "worksheet.docx"
    doc = Document()
    # Add Header
    header = doc.add_heading(title, level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add Sub-header
    sub_header_text = f"Date: {sub_title_info['date']}   |   Marks: {sub_title_info['marks']}   |   Topic: {sub_title_info['sub-title']}"
    sub_header = doc.add_paragraph(sub_header_text)
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("--------------------------------------------------") # Separator
    # Add AI Content
    doc.add_paragraph(content)
    # Add Footer (basic implementation)
    footer_section = doc.sections[0].footer
    if not footer_section.is_linked_to_previous: # Ensure footer isn't linked from previous section
         footer_p = footer_section.paragraphs[0] if footer_section.paragraphs else footer_section.add_paragraph()
         if footer_p:
            footer_p.text = "Generated by AI Worksheet Tool"
            footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.save(file_path) # Save the DOCX
    return file_path

def create_txt(content, title, sub_title_info):
    # ... (function code remains the same as before) ...
    print("Generating TXT...")
    file_path = "worksheet.txt"
    with open(file_path, "w", encoding="utf-8") as f:
        f.write(f"Title: {title}\n")
        f.write(f"Date: {sub_title_info['date']} | Marks: {sub_title_info['marks']}\n")
        f.write("-------------------------------------\n\n")
        f.write(content)
    return file_path

# --- Static File and Main Page Routes ---

# Serves index.html (main page) - Requires Login
@app.route('/')
@login_required
def serve_index():
    # Renders templates/index.html
    return render_template('index.html', user_email=current_user.email)

# Serves other static files (js, css) from the 'static' folder
@app.route('/static/<path:filename>')
def serve_static(filename):
     # Flask automatically handles security for send_from_directory
     return send_from_directory(app.static_folder, filename)

# --- AUTHENTICATION ROUTES ---

# --- Registration Route ---
@app.route('/register', methods=['GET', 'POST'])
def register():
    if current_user.is_authenticated:
        return redirect(url_for('serve_index')) # Redirect if already logged in

    if request.method == 'POST':
        # Get all form fields
        name = request.form.get('name')
        email = request.form.get('email')
        country_code = request.form.get('country_code')
        phone_number_main = request.form.get('phone_number_main')
        grade = request.form.get('grade')
        password = request.form.get('password')
        confirm_password = request.form.get('confirm_password')

        # --- Input Validation ---
        # Check required fields (Phone is required for standard registration)
        if not name or not email or not country_code or not phone_number_main or not password or not confirm_password:
            flash('Please fill out all required fields (Name, Email, Phone, Password).', 'warning')
            return redirect(url_for('register'))

        full_phone_number = country_code + phone_number_main.strip()

        # Validate Email Format
        email_regex = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        if not re.match(email_regex, email):
            flash('Invalid email format provided.', 'warning')
            return redirect(url_for('register'))

        # Validate Combined Phone Number Format
        phone_regex = r'^\+?\d{7,}$' # Basic check
        if not re.match(phone_regex, full_phone_number):
             flash('Invalid phone number format (must be at least 7 digits, starting with country code).', 'warning')
             return redirect(url_for('register'))

        # Validate Password Match
        if password != confirm_password:
            flash('Passwords do not match.', 'warning')
            return redirect(url_for('register'))

        # Validate Password Complexity
        if len(password) < 8: flash('Password requires at least 8 characters.', 'warning'); return redirect(url_for('register'))
        if not re.search(r'[A-Z]', password): flash('Password requires at least one uppercase letter.', 'warning'); return redirect(url_for('register'))
        if not re.search(r'[a-z]', password): flash('Password requires at least one lowercase letter.', 'warning'); return redirect(url_for('register'))
        if not re.search(r'\d', password): flash('Password requires at least one digit.', 'warning'); return redirect(url_for('register'))
        if not re.search(r'[!@#$%^&*()_+=\-\[\]{};\'\\:"|,.<>\/?~`]', password): flash('Password requires at least one special character.', 'warning'); return redirect(url_for('register'))
        if re.search(r'\s', password): flash('Password cannot contain spaces.', 'warning'); return redirect(url_for('register'))
        # --- End Validation ---

        # Check for uniqueness (Email and Phone Number)
        existing_email = User.query.filter_by(email=email).first()
        if existing_email:
            flash('An account with this email already exists.', 'warning')
            return redirect(url_for('register'))

        existing_phone = User.query.filter_by(phone_number=full_phone_number).first()
        if existing_phone:
             flash('An account with this phone number already exists.', 'warning')
             return redirect(url_for('register'))

        # Create new user instance
        new_user = User(name=name, email=email, phone_number=full_phone_number, grade=grade if grade else None)
        new_user.set_password(password) # Hash the password
        db.session.add(new_user) # Add to the session

        # Commit to database
        try:
            db.session.commit()
            flash('Registration successful! Please log in.', 'success')
            return redirect(url_for('login')) # Redirect to login page
        except Exception as e:
            db.session.rollback() # Rollback changes if commit fails
            flash(f'Database error during registration: {e}', 'danger')
            print(f"Database Error on registration commit: {e}") # Log detailed error
            return redirect(url_for('register'))

    # GET request: Display the registration form
    return render_template('register.html') # Renders templates/register.html

# --- Login Route (MODIFIED to accept Email or Phone) ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    if current_user.is_authenticated:
        return redirect(url_for('serve_index')) # Redirect if already logged in

    if request.method == 'POST':
        login_identifier = request.form.get('login_identifier', '').strip() # Get identifier and remove whitespace
        password = request.form.get('password')

        # --- UPDATED Validation Check ---
        if not login_identifier or not password:
            # More specific message
            flash('Please enter both your email/phone and password.', 'warning')
            return redirect(url_for('login'))
        # --- End Update ---

        # --- Try finding user by email OR phone (more robust) ---
        user = None
        # Check if it looks like an email
        if '@' in login_identifier and '.' in login_identifier:
            user = User.query.filter_by(email=login_identifier).first()
        else:
            # Assume it might be a phone number
            # Try matching exact input first (which includes country code if typed)
            user = User.query.filter_by(phone_number=login_identifier).first()

            # If not found AND input doesn't start with '+', try prepending a default
            # (Assumes stored numbers always start with '+countrycode' from registration)
            if user is None and not login_identifier.startswith('+'):
                 # Simple check: If it looks like a local number (e.g., >= 7 digits)
                 if len(login_identifier) >= 7:
                    # Example: Try adding +91 (adjust default country code if needed)
                    # This guess might fail if the user registered with a different code
                    potential_phone = "+91" + login_identifier # TODO: Make default code configurable?
                    print(f"Attempting login with assumed country code: {potential_phone}") # Debug print
                    user = User.query.filter_by(phone_number=potential_phone).first()
                    # Add more country code guesses here if necessary (e.g., +1)

        # ------------------------------------------

        # Check if a user was found and if the password is correct
        # For OAuth users without a password_hash, check_password will correctly return False
        if user is None or not user.check_password(password):
            flash('Invalid credentials. Please check your email/phone and password.', 'danger')
            return redirect(url_for('login'))

        # Log the user in
        login_user(user) # Flask-Login manages the session
        next_page = request.args.get('next') # For redirecting after @login_required
        return redirect(next_page or url_for('serve_index')) # Redirect to intended page or main app

    # GET request: Display the login form
    return render_template('login.html')

# --- Logout Route ---
@app.route('/logout')
@login_required
def logout():
    logout_user() # Clears session
    flash('You have been successfully logged out.', 'info')
    return redirect(url_for('login')) # Go back to login page

# --- GOOGLE OAUTH LOGIN ROUTES ---

@app.route('/login/google')
def login_google():
    redirect_uri = url_for('authorize_google', _external=True)
    if not app.config.get('GOOGLE_CLIENT_ID') or not app.config.get('GOOGLE_CLIENT_SECRET'):
         flash('Google login is not configured on the server.', 'danger'); print("ERROR: GOOGLE OAuth env vars missing.")
         return redirect(url_for('login'))
    # Authlib handles nonce internally
    return oauth.google.authorize_redirect(redirect_uri)

@app.route('/authorize/google')
def authorize_google():
    try:
        # Fetch the token first
        token = oauth.google.authorize_access_token()

        # Use userinfo endpoint instead of parsing id_token to avoid nonce issues
        resp = oauth.google.get('https://openidconnect.googleapis.com/v1/userinfo')
        resp.raise_for_status() # Raise exception for non-2xx status codes
        user_info = resp.json() # Get user info as dictionary

        if not user_info or not user_info.get('email'):
            flash('Could not fetch user information from Google.', 'danger')
            return redirect(url_for('login'))

        email = user_info.get('email')
        name = user_info.get('name') or user_info.get('given_name', email.split('@')[0]) # Use name or default

        # Find or create user based on email
        user = User.query.filter_by(email=email).first()
        if not user:
            # Create new user for Google login (phone/password/grade are null initially)
            user = User(email=email, name=name, phone_number=None, password_hash=None, grade=None)
            db.session.add(user)
            try:
                db.session.commit()
                flash('Account created via Google. Welcome!', 'success')
            except Exception as e:
                db.session.rollback(); flash(f'Error creating Google user: {e}', 'danger'); print(f"DB Error (Google OAuth): {e}")
                return redirect(url_for('login'))

        # Log the user in
        login_user(user)
        return redirect(url_for('serve_index')) # Redirect to main app page

    except Exception as e:
        # Catch errors during token fetching or userinfo request
        print(f"Google OAuth Error: {e}")
        import traceback
        traceback.print_exc() # Print full traceback for debugging OAuth issues
        flash(f'An error occurred during Google login. Please ensure Google credentials are correct. Error: {e}', 'danger')
        return redirect(url_for('login'))

# --- REMOVED FACEBOOK ROUTES ---

# ----------------------------------------

# --- Main AI Worksheet Generation API Endpoint ---
@app.route('/generate-worksheet', methods=['POST'])
@login_required # Protect this API
def generate_worksheet():
    try:
        # 1. Get data from JSON request
        data = request.json
        # Use current_user's grade if available and not provided in form
        grade = data.get('grade') or current_user.grade
        board = data.get('board')
        topic = data.get('topic')
        subtopic = data.get('subtopic')
        difficulty = data.get('difficulty')
        file_format = data.get('format')
        include_answers = data.get('answer_key', False)

        # 2. Validate essential inputs
        if not topic or not subtopic:
            return jsonify({"error": "Topic and Sub-Topic fields cannot be empty."}), 400
        if not grade: # Check if grade is still missing after potentially getting from profile
             return jsonify({"error": "Grade is required. Please specify in the form or update your profile."}), 400

        # Log request
        print(f"User '{current_user.email}' requested: Grade={grade}, Board={board}, Topic={topic}, SubTopic={subtopic}, Difficulty={difficulty}, Format={file_format}, Answers={include_answers}")

        # 3. Construct AI Prompt
        prompt = f"""
        You are an expert math teacher creating a worksheet.
        Generate 10 distinct math problems appropriate for the following criteria:
        - Grade: {grade}
        - Educational Board: {board} (Consider typical curriculum if known, otherwise general)
        - Topic: {topic}
        - Sub-Topic: {subtopic}
        - Difficulty Level: {difficulty} (Interpret easy/medium/hard appropriately for the grade level '{grade}')

        Instructions for Output:
        - Provide ONLY the list of 10 problems.
        - Do NOT include the answers directly with the problems.
        - Format the problems clearly, numbered 1. through 10.
        - IMPORTANT: Do not use any LaTeX math formatting (like $, \\frac, \\sqrt, etc.).
        - IMPORTANT: Do not use special symbols like square boxes (\\square).
        - If a space for an answer is needed within a problem, use exactly three underscores: ___
        """
        if include_answers:
            prompt += "\n\nFinally, after the 10 problems, create a separate section titled '--- ANSWER KEY ---' and provide the correct answers for all 10 problems, clearly numbered."

        # 4. Configure and Call Google Gemini API
        api_key = os.environ.get("GOOGLE_API_KEY")
        if not api_key:
             print("---!!! SERVER ERROR: GOOGLE_API_KEY environment variable not set. !!!---")
             return jsonify({"error": "Server configuration error. Please contact administrator."}), 500
        genai.configure(api_key=api_key)

        model = genai.GenerativeModel('gemini-flash-latest')
        safety_settings=[ # Define safety settings
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        ]
        response = model.generate_content(prompt, safety_settings=safety_settings)

        # 5. Process AI Response
        try:
            ai_content = response.text # Extract the generated text
        except ValueError: # Handle potential blocks more gracefully
            print("---!!! AI CONTENT GENERATION BLOCKED OR FAILED !!!---")
            print(f"Prompt Feedback: {response.prompt_feedback}")
            try: print(f"Candidates: {response.candidates}")
            except Exception: pass
            block_reason = response.prompt_feedback.block_reason if response.prompt_feedback else 'Unknown'
            # Give a more specific error if possible
            error_message = f"AI failed to generate content (Reason: {block_reason}). This may be due to safety filters or the specific topics requested. Please try simplifying your request."
            return jsonify({"error": error_message}), 500 # Use 500 for server-side AI failure
        except Exception as ai_err: # Catch other potential errors
             print(f"---!!! Error Processing AI Response: {ai_err} !!!---")
             return jsonify({"error": "An error occurred while processing the AI response."}), 500

        # 6. Clean AI Output
        ai_content = ai_content.replace(r"$\square$", "___").replace("\\square", "___").replace("$", "")

        print("--- Cleaned AI Generated Content ---")
        # print(ai_content) # Limit printing large content unless debugging
        print("--- End AI Content ---")

        # 7. Prepare Worksheet Metadata
        worksheet_title = f"Grade {grade} Math Worksheet: {topic}"
        worksheet_info = {
            "date": datetime.now().strftime("%Y-%m-%d"),
            "time": datetime.now().strftime("%H:%M:%S"),
            "marks": "___ / 50", # Placeholder
            "sub-title": subtopic
        }

        # 8. Generate and Send File
        file_path = ""
        mimetype = ""
        temp_pdf_for_image = None # Variable to hold temp PDF path for cleanup
        try: # Wrap file generation in try/finally for cleanup
            # --- File Generation Logic ---
            if file_format == 'pdf':
                file_path = create_pdf(ai_content, worksheet_title, worksheet_info)
                mimetype = 'application/pdf'
            elif file_format == 'docx':
                file_path = create_docx(ai_content, worksheet_title, worksheet_info)
                mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            elif file_format in ['png', 'jpg']:
                if not convert_from_path:
                     return jsonify({"error": f"Server error: {file_format.upper()} generation requires Poppler installation."}), 500
                # Generate PDF temporarily for conversion
                temp_pdf_for_image = create_pdf(ai_content, worksheet_title, worksheet_info, filename="temp_worksheet_for_image.pdf")
                try:
                    # Explicitly provide poppler_path if needed, adjust to your system
                    # poppler_bin_path = r"C:\path\to\your\poppler\bin" # Example for Windows
                    # images = convert_from_path(temp_pdf_for_image, poppler_path=poppler_bin_path)
                    images = convert_from_path(temp_pdf_for_image) # Assumes Poppler is in PATH
                except Exception as img_err:
                     print(f"---!!! PDF to {file_format.upper()} Conversion Error: {img_err} !!!---")
                     # Check if it's a Poppler path issue specifically
                     if "poppler" in str(img_err).lower():
                         error_detail = "Could not find Poppler. Ensure it's installed and its 'bin' directory is in your system's PATH environment variable."
                     else:
                         error_detail = str(img_err)
                     return jsonify({"error": f"Failed conversion to {file_format.upper()}. {error_detail}"}), 500

                if not images: return jsonify({"error": f"PDF to {file_format.upper()} conversion failed (no pages)."}), 500

                file_path = f"worksheet.{file_format}"
                if file_format == 'png':
                    images[0].save(file_path, "PNG")
                    mimetype = 'image/png'
                else: # JPG
                    images[0].convert('RGB').save(file_path, "JPEG") # Convert to RGB for JPG
                    mimetype = 'image/jpeg'
            else: # Default to TXT
                file_path = create_txt(ai_content, worksheet_title, worksheet_info)
                mimetype = 'text/plain'
            # --- End File Generation ---

            # Send the file as a download
            return send_file(file_path, as_attachment=True, mimetype=mimetype)

        finally:
             # Cleanup: Delete the temporary PDF if it was created
             if temp_pdf_for_image and os.path.exists(temp_pdf_for_image):
                 try:
                     os.remove(temp_pdf_for_image)
                     print(f"Cleaned up temporary file: {temp_pdf_for_image}")
                 except Exception as clean_err:
                     print(f"Error cleaning up temp file {temp_pdf_for_image}: {clean_err}")
             # Cleanup: Delete the final generated file after sending? Optional.
             # If you want to delete worksheet.pdf/docx/png/jpg after download:
             # if file_path and os.path.exists(file_path):
             #     try:
             #         # Note: send_file might need the file *during* sending.
             #         # Deleting immediately might cause issues. Consider background task.
             #         # os.remove(file_path)
             #         # print(f"Cleaned up output file: {file_path}")
             #         pass # Keep file for now
             #     except Exception as clean_err:
             #         print(f"Error cleaning up output file {file_path}: {clean_err}")


    # --- Global Error Handling for the Endpoint ---
    except Exception as e:
        print(f"---!!! UNEXPECTED ERROR in /generate-worksheet: {e} !!!---")
        import traceback
        traceback.print_exc() # Print full stack trace for debugging
        return jsonify({"error": "An internal server error occurred processing your request."}), 500

# --- Function to Create Database Tables ---
def create_database(app_instance):
    with app_instance.app_context():
        print("Initializing database...")
        db.create_all() # Creates tables if they don't exist based on models
        print("Database initialized.")

# --- Main Execution Block ---
if __name__ == '__main__':
    # Check for required OAuth environment variables
    missing_vars = [var for var in ['GOOGLE_CLIENT_ID', 'GOOGLE_CLIENT_SECRET'] if not app.config.get(var)]
    if missing_vars:
        print("\n---!!! WARNING: Google OAuth Environment Variables Missing !!!---")
        print("Google login will not work until you set:")
        for var in missing_vars: print(f"  - {var}")
        print("See instructions for setting environment variables.")
        print("-------------------------------------------------------\n")

    create_database(app) # Ensure database and tables are ready
    print("----------------------------------------------------")
    print("Flask server starting...")
    print(f"Database located at: {db_path}")
    print("Available Routes:")
    print("  - /                      (Main App - Requires Login)")
    print("  - /login                 (Login Page)")
    print("  - /register              (Registration Page)")
    print("  - /logout                (Logout Action - Requires Login)")
    print("  - /login/google          (Initiate Google Login)")
    print("  - /authorize/google      (Google Callback)")
    print("  - /generate-worksheet    (API Endpoint - Requires Login)")
    print("----------------------------------------------------")
    # Set debug=False for production deployment
    app.run(debug=True, port=5000) # Removed host='0.0.0.0' for standard local running

